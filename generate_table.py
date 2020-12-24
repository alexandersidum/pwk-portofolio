from docx import *
import re
from copy import deepcopy
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from XMLReader import XMLReader
from docx.shared import Pt
from docx.enum.style import WD_STYLE
import sys

def hapus_tabel_contoh(doc):
    # hapus tabel contoh pengisian
    table_contoh = doc.tables[-1]
    if table_contoh.cell(0,8).text == 'Diskripsi Evaluasi & Tindak lanjut perbaikan':
        table_contoh._element.getparent().remove(table_contoh._element)

    # hapus paragraph jelek
    paragraphs = doc.paragraphs
    for p in range(len(paragraphs)):
        if paragraphs[p].text == 'Portofolio penilaian & evaluasi proses dan hasil belajar {mhs_name=”Nama”; mhs_nrp=”NRP”}':
            [paragraphs[x]._element.getparent().remove(paragraphs[x]._element) for x in range(p-1,len(paragraphs))]

def tableNilaiCPL(doc):
    table = doc.tables[4]
    for row in range(len(table.rows)):
        if 'Peta CPL – CP MK' in table.cell(row,0).text:
            for col in range(len(table.columns)):
                str_peta_cpl = table.cell(row,col)
                if str_peta_cpl.tables:
                    return str_peta_cpl.tables[0]

def ambilBobotCPl(tabel):
    cpl_dict = {}
    for column in range(1,len(tabel.columns)):
        for row in range(1,len(tabel.rows)):
            if tabel.cell(row,column).text != '':
                cpl_key = tabel.cell(0,column).text
                cpl_value = tabel.cell(row,column).text
                cpl_dict[re.sub('CPL\s','',cpl_key)] = float(cpl_value)
    return cpl_dict

def CpmkCpl(tabel):
    cpmk_cpl = {}
    for row in range(1,len(tabel.rows)):
        cpl_isi=[]
        for column in range(1,len(tabel.columns)):
            if tabel.cell(row,column).text != '':
                cpmk_cpl_key = tabel.cell(row,0).text
                cpl = tabel.cell(0,column).text
                cpl_isi.append(re.sub('CPL\s','',cpl))
        cpmk_cpl[cpmk_cpl_key]=cpl_isi
    return cpmk_cpl

def ambilMinggu(table):
    minggu = []
    for row in range(1,len(table.rows)-1):
        minggu.append(table.cell(row,0).text)
    return minggu

def ambilCpmk(table):
    cpmk = []
    for row in range(1,len(table.rows)-1):
        cpmk.append(table.cell(row,1).text)
    return cpmk

def ambilBentukPenilaian(table):
    bentuk_penilaian = []
    for row in range(1,len(table.rows)-1):
        bentuk_penilaian.append(table.cell(row,2).text.splitlines()[0])
    return bentuk_penilaian

def ambilBobotCpmk(table):
    bobot = []
    for row in range(1,len(table.rows)-1):
        bobot.append(table.cell(row,3).text)
    return bobot

def ambilFailDesc(table):
    fail_desc = []
    for row in range(1,len(table.rows)-1):
        if len(table.cell(row,2).text.splitlines()) > 1:
            fail_desc.append(table.cell(row,2).text.splitlines()[-1])
        else:
            cpmk_ke = re.findall('CPMK\s\d',table.cell(row,1).text)[0]
            fail_desc.append(f'Tidak Lulus {cpmk_ke}')
    return fail_desc

def writeNormalStyle(cell,fill):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_bobot = cell.paragraphs[0]
    cell_bobot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_bobot.text = fill

def writeBoldStyle(cell,fill):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_paragraph.add_run(fill).bold = True

def writeCpmk(cell, fill):
    cell.text = fill.split('\n')[0]
    for line in fill.split('\n'):
        if 'CPMK' in line:
            pass
        else:
            cell.add_paragraph(line)

def writeCpl(cell, cpmk_cpl_dict, cpmk):
    cpl_cell = cell.paragraphs[0]
    cpl_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for key,value in cpmk_cpl_dict.items():
        if key in cpmk:
            cpl_cell.text = '\n'.join(value)
            break

def writeNiliaXBobot(cell,nilai,bobot):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_nilai_bobot = cell.paragraphs[0]
    cell_nilai_bobot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nilaixbobot = float(nilai)*(float(bobot)/100)
    cell_nilai_bobot.text = "{:.2f}".format(round(nilaixbobot, 2))

def writeKetercapaianCpl(cell,cpl_cell,cell_nilaixbobot,cpl_nilai_dict):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_ketercapaiancpl = cell.paragraphs[0]
    cell_ketercapaiancpl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nilaicpl = 0.0
    cpl = cpl_cell.text.split('\n')
    for key,value in cpl_nilai_dict.items():
        if key in cpl:
            nilaicpl += (value*float(cell_nilaixbobot.text))
    cell_ketercapaiancpl.text = str(round(nilaicpl,1))

def writeDesc(cell,nilai,cpmk,fail_desc):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_desc = cell.paragraphs[0]
    cell_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if float(nilai) > 50:
        cell_desc.text = 'Lulus {}'.format(cpmk.split("\n")[0])
    else:
        cell_desc.text = fail_desc

try:
    doc = Document(sys.argv[1])
    xml_data = XMLReader(sys.argv[2]).getData()

    hapus_tabel_contoh(doc)

    table_cpl = tableNilaiCPL(doc)
    bobot_cpl = ambilBobotCPl(table_cpl)
    cpmk_cpl = CpmkCpl(table_cpl)

    # cari table rae
    table_table = doc.tables
    for table in table_table:
        if table.cell(0,0).text == 'Mg ke\n(1)':
            rae_table = table

    cpmk = ambilCpmk(rae_table)
    mg_ke = ambilMinggu(rae_table)
    bentuk_penilaian = ambilBentukPenilaian(rae_table)
    bobot_cpmk = ambilBobotCpmk(rae_table)
    fail_desc = ambilFailDesc(rae_table)


    for x in xml_data:
        data_nilai = deepcopy(x[1:])

        template_doc = Document('template_tbl.docx')
        template_section = template_doc.sections[0]
        left, right, top, bottom, gutt = template_section.left_margin, template_section.right_margin, template_section.top_margin, template_section.bottom_margin, template_section.gutter
        new_width, new_height = template_section.page_width, template_section.page_height
        section = doc.add_section()
        section.left_margin, section.right_margin, section.top_margin, section.bottom_margin, section.gutter = left, right, top, bottom, gutt
        section.page_width , section.page_height = new_width, new_height
        keterangan = doc.add_paragraph().add_run('Portofolio penilaian & evaluasi proses dan hasil belajar {}'.format(data_nilai[0]))
        font = keterangan.font
        font.name = 'Britannic Bold'
        font.size = Pt(12)
        table = template_doc.tables[-1]
        new_tbl = deepcopy(table._tbl)
        paragraph =  doc.add_paragraph()
        paragraph._p.addnext(new_tbl)


        tbl = doc.tables[-1]
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'

        for i in range(len(mg_ke)):
            cell = tbl.add_row().cells
            writeBoldStyle(cell[0],mg_ke[i])
            writeCpl(cell[1],cpmk_cpl,cpmk[i])
            writeCpmk(cell[2],cpmk[i])
            writeNormalStyle(cell[3],bentuk_penilaian[i])
            writeNormalStyle(cell[4],bobot_cpmk[i])
            writeNormalStyle(cell[5],"{:.2f}".format(round(float(data_nilai[i+1]), 2)))
            writeNiliaXBobot(cell[6],data_nilai[i+1],bobot_cpmk[i])
            writeKetercapaianCpl(cell[7],cell[1],cell[6],bobot_cpl)
            writeDesc(cell[-1],data_nilai[i+1],cpmk[i],fail_desc[i])

    doc.save(sys.argv[3])

except IndexError:
    print(f'Usage : python {sys.argv[0]} [doc file] [xml file] [saved file]')

